from  pathlib import Path
import PyPDF2
import docx
from bs4 import BeautifulSoup
from typing import Dict,Any,List


def read_pdf(file_path):
    text = ""
    try:
        with open(file_path,"rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text+=extracted +"\n"

    except Exception as e:
        print(f"if we find in (file_path)in errors:{e}")
    return  text.strip()

import docx


def read_docx(file_path: str) -> str:
    text_chunks = []
    try:
        reader = docx.Document(file_path)

        # 1. Extract text from paragraphs
        for para in reader.paragraphs:
            cleaned = para.text.strip()
            if cleaned:
                text_chunks.append(cleaned)

        # 2. Extract text from tables
        for table in reader.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if row_text:
                    text_chunks.append(" | ".join(row_text))

        return "\n\n".join(text_chunks)

    except Exception as e:
        raise ValueError(f"Failed to read DOCX file: {str(e)}")

def read_txt(file_path):
    try:
        with open(file_path,"r",errors ="ignore",encoding="utf-8") as file:
            return file.read().strip()

    except Exception as e:
        print(f"if we find in (file_path)in errors:{e}")
        return ""

def read_html(file_path):
    text=""

    try:
        with open(file_path,"rb",errors ="ignore",encoding="utf-8") as file:
            html =file.read
            soup = BeautifulSoup(html,"html.parser")

        for script_for_style in soup(["script","style"]):
            script_for_style.decompose()
            return soup.get_text(separator="\n", strip=True)
    except Exception as e:
        print(f"if we find in (file_path)in errors:{e}")
        return ""


if __name__ == "__main__":

    file_path = read_pdf(r"C:\Users\hp\Downloads\Great.pdf")
    print(file_path[:1000])

    sample_pdf = read_txt(r"C:\Users\hp\Downloads\Document.txt")
    print(sample_pdf[:2000])
